#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""技术报告 markdown → DOCX（python-docx）与 PDF（reportlab）双渲染。

排版遵循 docx skill Profile A（正式报告）：黑体标题、宋体正文 12pt、
1.3 倍行距、正文首行缩进 2 字符、表头跨页重复、行禁拆。
"""
from __future__ import annotations

import hashlib
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "01_技术报告"
MD = REPORT_DIR / "技术报告_XH-202610.md"
OUT_DIR = REPORT_DIR / "交稿"
DOCX_OUT = OUT_DIR / "技术报告_XH-202610.docx"
PDF_OUT = OUT_DIR / "技术报告_XH-202610.pdf"
MATH_CACHE = OUT_DIR / "_math_cache"

IMG_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
ABS_RE = re.compile(r"^\*\*摘要\*\*[：:]\s*(.*)$", re.S)
KW_RE = re.compile(r"^\*\*关键词\*\*[：:]\s*(.*)$")
INLINE_MATH_RE = re.compile(r"(?<!\$)\$(?!\$)((?:\\.|[^$])+?)(?<!\$)\$(?!\$)")

# ---------------------------------------------------------------- 解析
Block = tuple  # (kind, payload)


def resolve_img(rel: str) -> Path:
    p = Path(rel)
    if not p.is_absolute():
        p = (REPORT_DIR / rel).resolve()
    return p


def normalize_latex(tex: str) -> str:
    """把 Markdown 公式改成 matplotlib mathtext 可渲染的子集。"""
    tex = re.sub(r"\s+", " ", tex.strip())
    tex = tex.replace(r"\arg\min", r"\mathrm{arg\,min}")
    tex = tex.replace(r"\arg\max", r"\mathrm{arg\,max}")
    tex = tex.replace(r"\bigl(", "(").replace(r"\bigr)", ")")
    tex = tex.replace(r"\bigl[", "[").replace(r"\bigr]", "]")
    tex = tex.replace(r"\bigl\{", r"\{").replace(r"\bigr\}", r"\}")
    tex = tex.replace(r"\bigl", "").replace(r"\bigr", "")
    tex = tex.replace(r"\left", "").replace(r"\right", "")
    tex = tex.replace(r"\mathrm{all\ OOF}", r"\mathrm{all\,OOF}")
    tex = tex.replace(r"\ge", r"\geq").replace(r"\le", r"\leq")
    tex = tex.replace(r"\times", r"\times")  # keep
    return tex


def render_math(tex: str, display: bool = True) -> Path | None:
    """用 matplotlib mathtext 把 LaTeX 渲成 PNG，供 DOCX/PDF 嵌入。"""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    tex_n = normalize_latex(tex)
    key = hashlib.md5(f"{int(display)}:{tex_n}".encode("utf-8")).hexdigest()[:14]
    MATH_CACHE.mkdir(parents=True, exist_ok=True)
    out = MATH_CACHE / f"m_{key}.png"
    if out.exists() and out.stat().st_size > 0:
        return out
    try:
        fig = plt.figure(figsize=(0.01, 0.01))
        fig.patch.set_facecolor("white")
        fig.text(
            0.5,
            0.5,
            f"${tex_n}$",
            ha="center",
            va="center",
            fontsize=13 if display else 11,
            color="black",
        )
        fig.savefig(
            out,
            dpi=220 if display else 200,
            bbox_inches="tight",
            pad_inches=0.1 if display else 0.04,
            facecolor="white",
            transparent=False,
        )
        plt.close(fig)
        return out
    except Exception as e:
        print("WARN math render failed:", e, "|", tex_n[:100], file=sys.stderr)
        try:
            plt.close("all")
        except Exception:
            pass
        return None


def math_png_width_cm(path: Path, max_cm: float = 15.5, min_cm: float = 4.0) -> float:
    from PIL import Image as PILImage

    with PILImage.open(path) as im:
        # 220 dpi → cm
        w_cm = im.width / 220.0 * 2.54
    return float(min(max_cm, max(min_cm, w_cm * 0.92)))


def extract_abstract_keywords(blocks):
    """从 **摘要** / **关键词** 段落或 > 引用块抽取。"""
    abs_text, kw_text = None, None
    for k, v in blocks:
        if k == "p":
            m = ABS_RE.match(v)
            if m:
                abs_text = m.group(1).strip()
                continue
            m = KW_RE.match(v)
            if m:
                kw_text = m.group(1).strip()
                continue
        elif k == "quote" and abs_text is None:
            abs_text = v
    if kw_text is None:
        kw_text = "运动想象；脑机接口；少样本个性化适配；迁移学习；异步交互协议"
    return abs_text or "", kw_text


def parse(md_text: str):
    lines = md_text.splitlines()
    blocks: list[Block] = []
    i, n = 0, len(lines)
    while i < n:
        ln = lines[i]
        s = ln.strip()
        if not s:
            i += 1
            continue
        if s.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", "\n".join(buf)))
            continue
        # 独立公式块 $$ ... $$
        if s == "$$":
            i += 1
            buf = []
            while i < n and lines[i].strip() != "$$":
                buf.append(lines[i].rstrip())
                i += 1
            if i < n:
                i += 1
            tex = " ".join(x.strip() for x in buf if x.strip())
            if tex:
                blocks.append(("math", tex))
            continue
        if s.startswith("$$") and s.endswith("$$") and len(s) > 4:
            blocks.append(("math", s[2:-2].strip()))
            i += 1
            continue
        if s.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            blocks.append(("table", rows))
            continue
        m_img = IMG_RE.match(s)
        if m_img:
            blocks.append(("img", m_img.group(2).strip()))
            i += 1
            continue
        if s.startswith("#### "):
            blocks.append(("h4", s[5:].strip()))
        elif s.startswith("### "):
            blocks.append(("h3", s[4:].strip()))
        elif s.startswith("## "):
            blocks.append(("h2", s[3:].strip()))
        elif s.startswith("# "):
            blocks.append(("h1", s[2:].strip()))
        elif s == "---":
            blocks.append(("hr", None))
        elif s.startswith("> "):
            blocks.append(("quote", s[2:].strip()))
        elif re.match(r"^\d+\.\s", s):
            blocks.append(("li_num", re.sub(r"^\d+\.\s+", "", s)))
        elif s.startswith("- "):
            blocks.append(("li_bul", s[2:].strip()))
        else:
            blocks.append(("p", s))
        i += 1
    return blocks


# ---------------------------------------------------------------- DOCX
def build_docx(blocks, out: Path):
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, Twips, Cm

    doc = Document()

    # 页面 A4 + 页边距
    for sec in doc.sections:
        sec.page_width, sec.page_height = Twips(11906), Twips(16838)
        sec.top_margin = sec.bottom_margin = Twips(1440)
        sec.left_margin = sec.right_margin = Twips(1600)

    def set_fonts(run, cn="宋体", en="Times New Roman", size=12, bold=False, color="000000"):
        run.font.name = en
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor.from_string(color)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), en)
        rfonts.set(qn("w:hAnsi"), en)
        rfonts.set(qn("w:eastAsia"), cn)

    def add_inline(p, text, size=12, bold_all=False, cn="宋体", allow_bold: bool = False):
        """解析代码/行内公式；正文默认不加粗（剥离 ** 标记），标题可 allow_bold。"""
        from docx.shared import Inches
        from PIL import Image as PILImage

        math_pat = re.compile(r"(\$\$.+?\$\$|\$(?:\\.|[^$])+?\$)")
        outer_pat = re.compile(r"(\*\*[^*]+?\*\*|`[^`]+`)")

        def emit_math(tex: str, display: bool = False):
            png = render_math(tex, display=display)
            if not png:
                r = p.add_run(tex)
                set_fonts(r, cn="Consolas", en="Consolas", size=size - 1)
                return
            if display:
                p.add_run().add_picture(str(png), width=Cm(math_png_width_cm(png, max_cm=14.0)))
                return
            with PILImage.open(png) as im:
                h_in = max(0.14, min(0.38, im.height / 200.0 * 0.95))
                w_in = im.width / im.height * h_in
            p.add_run().add_picture(str(png), width=Inches(w_in), height=Inches(h_in))

        def emit_text(s: str, bold: bool = False):
            if not s:
                return
            use_bold = bool(allow_bold and (bold or bold_all))
            for part in math_pat.split(s):
                if not part:
                    continue
                if part.startswith("$$") and part.endswith("$$") and len(part) > 4:
                    emit_math(part[2:-2], display=True)
                elif part.startswith("$") and part.endswith("$") and len(part) > 2:
                    emit_math(part[1:-1], display=False)
                else:
                    r = p.add_run(part)
                    set_fonts(r, cn=cn, size=size, bold=use_bold)

        for part in outer_pat.split(text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**") and len(part) >= 4:
                # 仅剥离 Markdown 加粗标记；正文字重不加粗
                emit_text(part[2:-2], bold=True)
            elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
                r = p.add_run(part[1:-1])
                set_fonts(r, cn="Consolas", en="Consolas", size=size - 1.5)
            else:
                emit_text(part, bold=False)

    def para(text, size=12, indent=480, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, cn="宋体"):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.3
        pf.space_before, pf.space_after = Pt(before), Pt(after)
        pf.alignment = align
        if indent:
            pf.first_line_indent = Twips(indent)
        add_inline(p, text, size=size, cn=cn)
        return p

    def heading(text, level, center=False):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.3
        pf.space_before, pf.space_after = Pt(14 if level == 1 else 10), Pt(8)
        # 一级标题与「摘要」居中；二、三级小节标题左对齐加粗
        if level == 1 or center or text.strip() in ("摘要", "摘  要", "摘 要"):
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hmap = {1: ("Heading 1", 16), 2: ("Heading 2", 15), 3: ("Heading 3", 14)}
        style, size = hmap.get(level, ("Heading 3", 13))
        p.style = doc.styles[style]
        add_inline(p, text, size=size, bold_all=True, cn="黑体", allow_bold=True)
        for r in p.runs:
            set_fonts(r, cn="黑体", size=size, bold=True)
        return p

    def shade(el, fill):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)
        el.append(shd)

    def table(rows):
        ncol = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=ncol)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 单元格边距
        tblpr = t._tbl.tblPr
        mar = OxmlElement("w:tblCellMar")
        for side, w in (("top", 40), ("bottom", 40), ("left", 80), ("right", 80)):
            e = OxmlElement(f"w:{side}")
            e.set(qn("w:w"), str(w))
            e.set(qn("w:type"), "dxa")
            mar.append(e)
        tblpr.append(mar)
        fsize = 10.5 if ncol <= 6 else 9
        for ri, row in enumerate(rows):
            trpr = t.rows[ri]._tr.get_or_add_trPr()
            cant = OxmlElement("w:cantSplit")
            trpr.append(cant)
            if ri == 0:
                th = OxmlElement("w:tblHeader")
                trpr.append(th)
            for ci in range(ncol):
                cell = t.cell(ri, ci)
                text = row[ci] if ci < len(row) else ""
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
                add_inline(p, text, size=fsize, bold_all=False, allow_bold=False)
                if ri == 0:
                    shade(cell._tc.get_or_add_tcPr(), "D9E2F3")

    # ---- 封面（简单标题页）----
    title = next(v for k, v in blocks if k == "h1")
    meta = next(v for k, v in blocks if k == "p" and v.startswith("**题目编号"))
    abs_text, kw_text = extract_abstract_keywords(blocks)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(200)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_fonts(r, cn="黑体", size=24, bold=True)
    doc.add_paragraph()
    for mtext in ("基于运动想象的脑-机交互算法研究与系统实现", meta.strip("*")):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(mtext)
        set_fonts(r, cn="宋体", size=14)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- 摘要 ----
    heading("摘要", 2, center=True)
    if abs_text:
        para(abs_text, size=12, indent=480)
    para(f"**关键词：**{kw_text}", indent=480)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- 目录（域，需在 Word 中更新）----
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录")
    set_fonts(r, cn="黑体", size=16, bold=True)
    fldp = doc.add_paragraph()
    r1 = fldp.add_run()
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin"); fc1.set(qn("w:dirty"), "true")
    r1._r.append(fc1)
    r2 = fldp.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = r' TOC \o "1-3" \h \z \u '
    r2._r.append(it)
    r3 = fldp.add_run()
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "separate")
    r3._r.append(fc2)
    r4 = fldp.add_run("（打开后按 F9 更新目录）")
    set_fonts(r4, cn="宋体", size=10.5)
    r5 = fldp.add_run()
    fc3 = OxmlElement("w:fldChar"); fc3.set(qn("w:fldCharType"), "end")
    r5._r.append(fc3)
    hint = doc.add_paragraph()
    hint.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hint.add_run("提示：打开文档后按 Ctrl+A → F9（或右键目录 → 更新域）刷新目录页码。")
    set_fonts(r, cn="宋体", size=9)
    r.font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor.from_string("808080")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- 正文 ----
    hr_count = 0
    seen_body = False
    for k, v in blocks:
        if k == "hr":
            hr_count += 1
            if hr_count >= 2:
                seen_body = True
            continue
        if not seen_body:
            continue
        if k == "h1":
            continue  # 封面已用
        if k == "h2":
            heading(v, 1)
        elif k == "h3":
            heading(v, 2)
        elif k == "h4":
            heading(v, 3)
        elif k == "img":
            path = resolve_img(v)
            if not path.exists():
                print("WARN missing image:", path, file=sys.stderr)
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            p.add_run().add_picture(str(path), width=Cm(14.5))
        elif k == "math":
            path = render_math(v, display=True)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            if path and path.exists():
                p.add_run().add_picture(str(path), width=Cm(math_png_width_cm(path)))
            else:
                add_inline(p, v, size=10.5, cn="Consolas")
        elif k == "p":
            if v.startswith("**题目编号"):
                continue
            if ABS_RE.match(v) or KW_RE.match(v):
                continue
            # 图注居中、无首行缩进
            if re.match(r"^\*\*图\s*\d+|^\*\*附图", v):
                para(v, size=10.5, indent=0, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=8)
            else:
                para(v)
        elif k == "quote":
            para(v, size=10.5, indent=480)
        elif k == "code":
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing = 1.1
            pf.left_indent = Twips(480)
            shade(p._p.get_or_add_pPr(), "F2F2F2")
            for seg in v.split("\n"):
                r = p.add_run(seg)
                set_fonts(r, cn="宋体", en="Consolas", size=9)
                r.add_break()
        elif k == "table":
            table(v)
            doc.add_paragraph().paragraph_format.space_after = Pt(3)
        elif k == "li_num":
            p = para(v, indent=0)
            p.paragraph_format.left_indent = Twips(480)
            add_inline(p, "· ", size=12)
            p.runs.clear() if False else None
        elif k == "li_bul":
            p = para(v, indent=0)
            p.paragraph_format.left_indent = Twips(480)

    # 页脚页码
    for sec in doc.sections:
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE \\* arabic \\* MERGEFORMAT")
        fp._p.append(fld)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print("docx saved:", out)


# ---------------------------------------------------------------- PDF
def build_pdf(blocks, out: Path):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        BaseDocTemplate, Frame, PageTemplate, Paragraph, Preformatted,
        Spacer, Table, TableStyle, PageBreak, Image as RLImage,
    )
    from reportlab.platypus.tableofcontents import TableOfContents
    from reportlab.lib.utils import ImageReader

    FONT_DIR = Path(r"C:\Windows\Fonts")
    pdfmetrics.registerFont(TTFont("SimSun", str(FONT_DIR / "simsun.ttc"), subfontIndex=0))
    pdfmetrics.registerFont(TTFont("SimHei", str(FONT_DIR / "simhei.ttf")))
    pdfmetrics.registerFontFamily("SimSun", normal="SimSun", bold="SimHei", italic="SimSun", boldItalic="SimHei")

    BLACK = colors.HexColor("#000000")
    HEAD_BG = colors.HexColor("#D9E2F3")
    CODE_BG = colors.HexColor("#F2F2F2")
    GRAY = colors.HexColor("#606060")

    def st(name, **kw):
        base = dict(fontName="SimSun", fontSize=12, leading=12 * 1.3, textColor=BLACK,
                    alignment=TA_JUSTIFY, spaceAfter=6, wordWrap="CJK")
        base.update(kw)
        return ParagraphStyle(name, **base)

    S = {
        "h1": st("h1", fontName="SimHei", fontSize=16, leading=16 * 1.3, alignment=TA_CENTER,
                 spaceBefore=14, spaceAfter=8, keepWithNext=1),
        "h2": st("h2", fontName="SimHei", fontSize=15, leading=15 * 1.3, spaceBefore=10, spaceAfter=8,
                 alignment=TA_LEFT, keepWithNext=1),
        "h3": st("h3", fontName="SimHei", fontSize=13, leading=13 * 1.3, spaceBefore=8, spaceAfter=5,
                 alignment=TA_LEFT, keepWithNext=1),
        "abstract_title": st("abstract_title", fontName="SimHei", fontSize=15, leading=15 * 1.3,
                             alignment=TA_CENTER, spaceBefore=10, spaceAfter=8, keepWithNext=1),
        "body": st("body", firstLineIndent=24),
        "quote": st("quote", fontSize=10.5, leading=10.5 * 1.3, firstLineIndent=21, textColor=BLACK),
        "code": st("code", fontName="Courier", fontSize=9, leading=11, alignment=TA_LEFT),
        "code_cjk": st("code_cjk", fontName="SimSun", fontSize=9.5, leading=13, alignment=TA_LEFT),
        "li": st("li", leftIndent=24, firstLineIndent=0),
        "cell": st("cell", fontSize=9, leading=11.5, alignment=TA_LEFT, spaceAfter=0),
        "cellh": st("cellh", fontName="SimSun", fontSize=9, leading=11.5, alignment=TA_CENTER, spaceAfter=0),
        "cover_title": st("cover_title", fontName="SimHei", fontSize=24, leading=34, alignment=TA_CENTER),
        "cover_meta": st("cover_meta", fontSize=14, leading=20, alignment=TA_CENTER),
    }

    def esc(t):
        return (t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                 .replace("\u2212", "-").replace("\u2010", "-"))

    def latex_to_plain(tex: str) -> str:
        s = normalize_latex(tex)
        repl = [
            (r"\theta", "θ"), (r"\tau", "τ"), (r"\pi", "π"), (r"\eta", "η"),
            (r"\mu", "μ"), (r"\sigma", "σ"), (r"\ell", "ℓ"),
            (r"\Delta", "Δ"), (r"\nabla", "∇"),
            (r"\mathcal{L}", "ℒ"), (r"\mathcal{B}", "ℬ"), (r"\mathcal{J}", "𝒥"),
            (r"\mathbb{R}", "ℝ"), (r"\times", "×"), (r"\ldots", "…"),
            (r"\geq", "≥"), (r"\leq", "≤"), (r"\in", "∈"), (r"\mid", "∣"),
            (r"\cdot", "·"), (r"\approx", "≈"), (r"\sum", "∑"),
            (r"\quad", " "), (r"\qquad", "  "), (r"\,", " "), (r"\ ", " "),
            (r"\log", "log"), (r"\exp", "exp"), (r"\max", "max"), (r"\min", "min"),
            (r"\mathrm{arg\,min}", "argmin"), (r"\mathrm{arg\,max}", "argmax"),
            (r"\mathrm{mode}", "mode"), (r"\mathrm{val}", "val"),
            (r"\mathrm{test}", "test"), (r"\mathrm{Acc}", "Acc"),
            (r"\tilde", "~"), (r"\hat", "^"),
        ]
        for a, b in repl:
            s = s.replace(a, b)
        s = re.sub(r"\\[a-zA-Z]+", "", s)
        s = s.replace("{", "").replace("}", "").replace("^", "")
        return s

    def inline(t):
        # 行内公式：不用 Paragraph <img>（CJK 换行会崩）；转成可读符号。
        # 独立 $$ 公式块另以 PNG 嵌入（见 k=="math"）。
        parts = []
        last = 0
        for m in INLINE_MATH_RE.finditer(t):
            parts.append(("t", t[last:m.start()]))
            parts.append(("m", m.group(1)))
            last = m.end()
        parts.append(("t", t[last:]))

        out = []
        for kind, payload in parts:
            if kind == "t":
                if not payload:
                    continue
                s = esc(payload)
                # 正文不加粗：仅剥离 Markdown ** 标记
                s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)

                def _code(m):
                    s2 = m.group(1)
                    if any("\u4e00" <= ch <= "\u9fff" for ch in s2):
                        return s2
                    return f'<font face="Courier" size="10">{s2}</font>'

                s = re.sub(r"`([^`]+)`", _code, s)
                out.append(s)
            else:
                out.append(f"<i>{esc(latex_to_plain(payload))}</i>")
        return "".join(out)

    story: list = []
    toc = TableOfContents()
    toc.levelStyles = [
        st("toc1", fontName="SimHei", fontSize=12, leading=18, alignment=TA_LEFT, leftIndent=6),
        st("toc2", fontSize=10.5, leading=15, alignment=TA_LEFT, leftIndent=24),
    ]

    # 封面
    title = next(v for k, v in blocks if k == "h1")
    meta = next(v for k, v in blocks if k == "p" and v.startswith("**题目编号")).strip("*")
    abs_text, kw_text = extract_abstract_keywords(blocks)
    meta_parts = meta.rsplit("｜", 1)
    meta_lines = [p_.strip() for p_ in meta_parts if p_.strip()] if len(meta_parts) == 2 else [meta]
    story += [Spacer(1, 55 * mm), Paragraph(esc(title), S["cover_title"]), Spacer(1, 8 * mm)]
    story.append(Paragraph(esc("基于运动想象的脑-机交互算法研究与系统实现"), S["cover_meta"]))
    story.append(Spacer(1, 3 * mm))
    for ml in meta_lines:
        story.append(Paragraph(esc(ml), S["cover_meta"]))
    story.append(PageBreak())

    # 摘要
    story.append(Paragraph("摘  要", S["abstract_title"]))
    if abs_text:
        story.append(Paragraph(inline(abs_text), S["body"]))
    story.append(Paragraph(inline(f"**关键词：**{kw_text}"), S["body"]))
    story.append(PageBreak())

    # 目录
    story.append(Paragraph("目  录", S["h1"]))
    story.append(Spacer(1, 4 * mm))
    story.append(toc)
    story.append(PageBreak())

    # 正文（带 TOC 收集）：第二个分隔线之后
    hr_count = 0
    seen_body = False

    class Doc(BaseDocTemplate):
        def afterFlowable(self, fl):
            if isinstance(fl, Paragraph) and getattr(fl, "_toc_entry", None):
                level, text = fl._toc_entry
                self.notify("TOCEntry", (level, text, self.page))

    def h(text, level):
        p = Paragraph(inline(text), S[{1: "h1", 2: "h2", 3: "h3"}[level]])
        p._toc_entry = (level - 1, re.sub(r"\*\*|`", "", text))
        return p

    def make_img(rel: str):
        path = resolve_img(rel)
        if not path.exists():
            print("WARN missing image:", path, file=sys.stderr)
            return None
        ir = ImageReader(str(path))
        iw, ih = ir.getSize()
        max_w = 150 * mm
        max_h = 170 * mm
        w = max_w
        h = w * ih / iw
        if h > max_h:
            h = max_h
            w = h * iw / ih
        return RLImage(str(path), width=w, height=h)

    for k, v in blocks:
        if k == "hr":
            hr_count += 1
            if hr_count >= 2:
                seen_body = True
            continue
        if not seen_body:
            continue
        if k == "h1":
            continue
        if k == "h2":
            story.append(h(v, 1))
        elif k == "h3":
            story.append(h(v, 2))
        elif k == "h4":
            story.append(h(v, 3))
        elif k == "img":
            img = make_img(v)
            if img is not None:
                story.append(Spacer(1, 2 * mm))
                story.append(img)
                story.append(Spacer(1, 2 * mm))
        elif k == "math":
            png = render_math(v, display=True)
            if png and png.exists():
                from PIL import Image as PILImage

                with PILImage.open(png) as im:
                    max_w = 155 * mm
                    w = min(max_w, im.width / 220.0 * 25.4 * mm * 0.92)
                    img_h = im.height / im.width * w
                story.append(Spacer(1, 2 * mm))
                story.append(RLImage(str(png), width=w, height=img_h))
                story.append(Spacer(1, 2 * mm))
            else:
                story.append(Paragraph(esc(v), S["code"]))
        elif k == "p":
            if v.startswith("**题目编号"):
                continue
            if ABS_RE.match(v) or KW_RE.match(v):
                continue
            if re.match(r"^\*\*图\s*\d+|^\*\*附图", v):
                story.append(Paragraph(inline(v), st(
                    "caption", fontSize=10, leading=13, alignment=TA_CENTER,
                    firstLineIndent=0, spaceBefore=2, spaceAfter=8)))
            else:
                story.append(Paragraph(inline(v), S["body"]))
        elif k == "quote":
            story.append(Paragraph(inline(v), S["quote"]))
        elif k == "code":
            has_cjk = any("\u4e00" <= ch <= "\u9fff" for ch in v)
            style = S["code_cjk"] if has_cjk else S["code"]
            story.append(Preformatted(v, style, maxLineLength=90))
            story.append(Spacer(1, 3 * mm))
        elif k == "table":
            ncol = max(len(r) for r in v)
            data = []
            for ri, row in enumerate(v):
                data.append([Paragraph(inline(row[ci] if ci < len(row) else ""), S["cellh" if ri == 0 else "cell"])
                             for ci in range(ncol)])
            t = Table(data, colWidths=None, repeatRows=1)
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#909090")),
                ("BACKGROUND", (0, 0), (-1, 0), HEAD_BG),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F7F5")]),
            ]))
            story += [t, Spacer(1, 3 * mm)]
        elif k in ("li_num", "li_bul"):
            story.append(Paragraph("• " + inline(v), S["li"]))

    def on_page(canv, doc_):
        canv.saveState()
        canv.setFont("SimSun", 9)
        canv.drawCentredString(A4[0] / 2, 12 * mm, str(canv.getPageNumber()))
        canv.restoreState()

    out.parent.mkdir(parents=True, exist_ok=True)
    doc_ = Doc(str(out), pagesize=A4,
               leftMargin=25 * mm, rightMargin=25 * mm, topMargin=22 * mm, bottomMargin=22 * mm,
               title="基于运动想象的脑-机交互算法研究与系统实现技术报告")
    frame = Frame(doc_.leftMargin, doc_.bottomMargin, doc_.width, doc_.height, id="main")
    doc_.addPageTemplates([PageTemplate(id="all", frames=[frame], onPage=on_page)])
    doc_.multiBuild(story)  # 多轮构建以生成目录条目
    print("pdf saved:", out)


if __name__ == "__main__":
    # 默认导出交稿定稿；可用环境变量 MD_SRC 覆盖
    import os

    md_src = os.environ.get("MD_SRC")
    if md_src:
        MD = Path(md_src)
    elif (OUT_DIR / "技术报告_XH-202610_v4.md").exists():
        MD = OUT_DIR / "技术报告_XH-202610_v4.md"
    _old_resolve = resolve_img

    def resolve_img(rel: str) -> Path:  # noqa: F811
        path = _old_resolve(rel)
        if path.exists():
            return path
        alt = OUT_DIR / rel
        return alt if alt.exists() else path

    blocks = parse(MD.read_text(encoding="utf-8"))
    print("source:", MD)
    build_docx(blocks, DOCX_OUT)
    build_pdf(blocks, PDF_OUT)
